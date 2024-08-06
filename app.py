import os
import streamlit as st
import cv2
import numpy as np
from PIL import Image
import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
from io import BytesIO

# Load environment variables from .env file
load_dotenv()

# Configure Gemini API
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)

generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-pro",
    generation_config=generation_config,
)

def preprocess_image(image, scale=1.0):
    """Convert image to grayscale, apply resizing based on scale, and preprocessing steps."""
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    # Resize image based on scale
    width = int(gray.shape[1] * scale)
    height = int(gray.shape[0] * scale)
    dim = (width, height)
    resized = cv2.resize(gray, dim, interpolation=cv2.INTER_AREA)
    blurred = cv2.GaussianBlur(resized, (5, 5), 0)
    edges = cv2.Canny(blurred, 50, 150)
    kernel = np.ones((5, 5), np.uint8)
    morphed = cv2.morphologyEx(edges, cv2.MORPH_CLOSE, kernel)
    return morphed

def count_and_detail_stamps(image, min_area=500):
    """Count and detail stamps in the image."""
    morphed = preprocess_image(image)
    contours, _ = cv2.findContours(morphed, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    
    stamp_details = []
    unique_bounding_boxes = set()  # To track unique bounding boxes

    for cnt in contours:
        area = cv2.contourArea(cnt)
        if area > min_area:
            x, y, w, h = cv2.boundingRect(cnt)
            bounding_box = (x, y, w, h)
            if bounding_box not in unique_bounding_boxes:
                unique_bounding_boxes.add(bounding_box)
                stamp_details.append({
                    'area': area,
                    'bounding_box': bounding_box,
                    'position': (x, y),
                    'size': (w, h),
                    'aspect_ratio': w / h,
                    'centroid': (x + w // 2, y + h // 2)
                })
    
    return len(stamp_details), stamp_details

def draw_bounding_boxes(image, stamp_details):
    """Draw bounding boxes around detected stamps and annotate them with details."""
    for detail in stamp_details:
        x, y, w, h = detail['bounding_box']
        # Draw bounding box
        cv2.rectangle(image, (x, y), (x+w, y+h), (0, 255, 0), 2)
        
        # Prepare text with details
        text = (
            f"Area: {detail['area']} px\n"
            f"Pos: ({detail['position'][0]}, {detail['position'][1]})\n"
            f"Size: {detail['size'][0]}x{detail['size'][1]} px\n"
            f"Aspect Ratio: {detail['aspect_ratio']:.2f}\n"
            f"Centroid: ({detail['centroid'][0]}, {detail['centroid'][1]})"
        )
        
        # Set the font and text position
        font = cv2.FONT_HERSHEY_SIMPLEX
        font_scale = 0.5
        font_color = (0, 255, 0)
        line_type = 1
        
        # Calculate text size
        (text_width, text_height), _ = cv2.getTextSize(text, font, font_scale, line_type)
        
        # Draw background rectangle for text
        cv2.rectangle(image, (x, y - text_height - 10), (x + text_width, y), (0, 0, 0), -1)
        
        # Add text to image
        cv2.putText(image, text, (x, y - 5), font, font_scale, font_color, line_type)
    
    return image

def generate_doc(stamp_count, stamp_details, analysis_response):
    """Generate a DOC file with analysis results."""
    doc = Document()
    doc.add_heading('Stamp Analysis Report', 0)
    
    doc.add_paragraph(f'Number of stamps detected: {stamp_count}')
    doc.add_paragraph('Details of detected stamps:')
    
    for i, detail in enumerate(stamp_details, 1):
        doc.add_paragraph(f'Stamp {i}:')
        doc.add_paragraph(f'  - Area: {detail["area"]}')
        doc.add_paragraph(f'  - Position: {detail["position"]}')
        doc.add_paragraph(f'  - Size: {detail["size"]}')
        doc.add_paragraph(f'  - Aspect Ratio: {detail["aspect_ratio"]:.2f}')
        doc.add_paragraph(f'  - Centroid: {detail["centroid"]}')
    
    doc.add_paragraph('Analysis Result:')
    doc.add_paragraph(analysis_response)
    
    # Save document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def analyze_image(image, min_area=500, scale=1.0):
    """Analyze the image and generate a description using the Gemini model."""
    stamp_count, stamp_details = count_and_detail_stamps(image, min_area)
    
    annotated_image = draw_bounding_boxes(image.copy(), stamp_details)
    
    _, img_encoded = cv2.imencode('.jpg', annotated_image)
    img_bytes = img_encoded.tobytes()
    
    description = f"Analyze this image and provide details about the {stamp_count} stamps detected."
    content_parts = [
        description,
        {"mime_type": "image/jpeg", "data": img_bytes}
    ]
    
    try:
        response = model.generate_content(content_parts)
        return stamp_count, stamp_details, response.text
    except Exception as e:
        st.error(f"Error generating analysis: {e}")
        return stamp_count, stamp_details, "Analysis failed."

# Streamlit UI
st.title("Philatelic Stamp Recognition and Analysis Platform")
uploaded_file = st.file_uploader("Choose an image...", type=["jpg", "jpeg", "png"])

if uploaded_file is not None:
    image = Image.open(uploaded_file)
    image_np = np.array(image)
    
    st.image(image, caption='Uploaded Image', use_column_width=True)

    # Fixed values for minimum area and scale
    min_area = 1200
    scale = 1.0

    if st.button("Analyze Image and Generate Summary"):
        with st.spinner("Analyzing..."):
            stamp_count, stamp_details, analysis_response = analyze_image(image_np, min_area, scale)
        
        st.write(f"Number of stamps detected: {stamp_count}")
        st.write("Analysis Result:")
        st.write(analysis_response)

        # Generate and provide download link for the DOC file
        doc_buffer = generate_doc(stamp_count, stamp_details, analysis_response)
        st.download_button(
            label="Download Report",
            data=doc_buffer,
            file_name="stamp_analysis_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
